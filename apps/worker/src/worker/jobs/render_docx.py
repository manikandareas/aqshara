from __future__ import annotations

from docx import Document


def render_document_value_to_docx_bytes(*, title: str, value: list[dict[str, object]]) -> bytes:
    document = Document()
    if title.strip():
        document.add_heading(title.strip(), level=0)

    for block in value:
        block_type = block["type"]
        if block_type == "heading":
            level = int(block.get("level", 1))
            text = "".join(child["text"] for child in block["children"])
            document.add_heading(text, level=min(level, 3))
        elif block_type == "paragraph":
            text = "".join(child["text"] for child in block["children"])
            document.add_paragraph(text)
        elif block_type == "bullet-list":
            for item in block["children"]:
                text = "".join(child["text"] for child in item["children"])
                document.add_paragraph(text, style="List Bullet")

    buffer = b""
    from io import BytesIO

    io_buffer = BytesIO()
    document.save(io_buffer)
    buffer = io_buffer.getvalue()
    return buffer
